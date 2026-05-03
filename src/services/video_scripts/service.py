"""Video script service — Phase 1+2.

- Template management (CRUD)
- Multi-variant AI script generation per video (parallel A/B/C)
- Scene-level edit (caption / b_roll_brief / duration / hashtags)
- DOCX export
"""

from __future__ import annotations

import json
import os
import sqlite3
import time
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from src.core.logger import get_logger
from src.services.ai.router import AIRouter, get_router
from src.services.video_scripts.prompts import (
    SCRIPT_EXTRACT_INSTRUCTIONS,
    SCRIPT_EXTRACT_SCHEMA,
    SCRIPT_SYSTEM_PROMPT,
    build_script_main_prompt,
)

logger = get_logger("service.video_scripts")

DEFAULT_DB_PATH = Path("data/ops_workbench.db")

# ----------------------------------------------------------------------
# Default 5-scene template (caption-only sticker pack TikTok ~33s)
# ----------------------------------------------------------------------

DEFAULT_TEMPLATE_NAME = "贴纸 5 段式标准模板（30-35s）"
DEFAULT_TEMPLATE_BLUEPRINT = [
    {
        "kind": "hook",
        "label": "Hook 开场",
        "duration_s": 3,
        "intent": "1 秒钩住注意力，制造好奇心或情绪共鸣",
        "visual_brief": "单张最有梗/最炸的贴纸全屏闪现 + 大字幕 POV/Wait...",
    },
    {
        "kind": "pack_showcase",
        "label": "完整卡包展示",
        "duration_s": 7,
        "intent": "让买家看到一包能拿到多少 / 哪些设计",
        "visual_brief": "桌面平铺整包 → 镜头从上向下扫过 → 字幕标注数量",
    },
    {
        "kind": "production",
        "label": "生成环节展示",
        "duration_s": 10,
        "intent": "建立信任 + 展示工艺 / AI 设计感",
        "visual_brief": "屏录 AI 设计稿 → 实拍打印切割 → 包装入袋（快剪）",
    },
    {
        "kind": "use_demo",
        "label": "成品使用展示",
        "duration_s": 10,
        "intent": "场景代入，激发购买动机",
        "visual_brief": "贴 laptop / 水杯 / 手账 / 手机壳 3-4 个快切场景",
    },
    {
        "kind": "cta",
        "label": "CTA 收尾",
        "duration_s": 5,
        "intent": "明确转化指令 + 包装定格",
        "visual_brief": "包装袋定格 + 大字幕 'Tap to shop ↑' + 店铺 logo 角标",
    },
]


def _open_db(db_path: Path = DEFAULT_DB_PATH) -> sqlite3.Connection:
    conn = sqlite3.connect(str(db_path), timeout=30, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def _now() -> int:
    return int(time.time())


class VideoScriptService:
    def __init__(
        self,
        router: Optional[AIRouter] = None,
        db_path: Path = DEFAULT_DB_PATH,
    ) -> None:
        self.router = router or get_router()
        self.db_path = db_path
        self._ensure_default_template()

    # ------------------------------------------------------------------
    # Template management
    # ------------------------------------------------------------------

    def _ensure_default_template(self) -> None:
        with _open_db(self.db_path) as conn:
            row = conn.execute(
                "SELECT id FROM tk_video_script_templates WHERE is_default = 1 LIMIT 1",
            ).fetchone()
            if row:
                return
            n = _now()
            conn.execute(
                """INSERT INTO tk_video_script_templates
                    (name, description, music_style, scene_blueprint,
                     is_default, is_archived, created_at, updated_at)
                   VALUES (?, ?, ?, ?, 1, 0, ?, ?)""",
                (
                    DEFAULT_TEMPLATE_NAME,
                    "默认模板：5 段式贴纸视频（Hook → 卡包 → 生成 → 使用 → CTA），约 33 秒。",
                    "soft uplifting vlog 95-110 bpm",
                    json.dumps(DEFAULT_TEMPLATE_BLUEPRINT, ensure_ascii=False),
                    n, n,
                ),
            )
            conn.commit()

    def list_templates(self, *, include_archived: bool = False) -> list[dict]:
        with _open_db(self.db_path) as conn:
            rows = conn.execute(
                f"""SELECT id, name, description, music_style, is_default, is_archived,
                          scene_blueprint, created_at, updated_at
                     FROM tk_video_script_templates
                    {"" if include_archived else "WHERE is_archived = 0"}
                    ORDER BY is_default DESC, id DESC""",
            ).fetchall()
        out = []
        for r in rows:
            d = dict(r)
            try:
                d["scene_blueprint"] = json.loads(d.get("scene_blueprint") or "[]")
            except Exception:
                d["scene_blueprint"] = []
            d["total_duration_s"] = sum(
                int(s.get("duration_s") or 0) for s in d["scene_blueprint"]
            )
            d["scene_count"] = len(d["scene_blueprint"])
            out.append(d)
        return out

    def get_template(self, template_id: int) -> Optional[dict]:
        with _open_db(self.db_path) as conn:
            r = conn.execute(
                """SELECT * FROM tk_video_script_templates WHERE id = ?""",
                (template_id,),
            ).fetchone()
        if not r:
            return None
        d = dict(r)
        try:
            d["scene_blueprint"] = json.loads(d.get("scene_blueprint") or "[]")
        except Exception:
            d["scene_blueprint"] = []
        return d

    def upsert_template(
        self,
        *,
        template_id: Optional[int] = None,
        name: str,
        description: str = "",
        music_style: str = "",
        scene_blueprint: list[dict],
        is_default: bool = False,
    ) -> int:
        n = _now()
        bp_json = json.dumps(scene_blueprint, ensure_ascii=False)
        with _open_db(self.db_path) as conn:
            if is_default:
                conn.execute(
                    "UPDATE tk_video_script_templates SET is_default = 0",
                )
            if template_id:
                conn.execute(
                    """UPDATE tk_video_script_templates
                          SET name=?, description=?, music_style=?,
                              scene_blueprint=?, is_default=?, updated_at=?
                        WHERE id=?""",
                    (name, description, music_style, bp_json,
                     1 if is_default else 0, n, template_id),
                )
                conn.commit()
                return template_id
            cur = conn.execute(
                """INSERT INTO tk_video_script_templates
                    (name, description, music_style, scene_blueprint,
                     is_default, is_archived, created_at, updated_at)
                   VALUES (?, ?, ?, ?, ?, 0, ?, ?)""",
                (name, description, music_style, bp_json,
                 1 if is_default else 0, n, n),
            )
            conn.commit()
            return cur.lastrowid

    def delete_template(self, template_id: int) -> bool:
        with _open_db(self.db_path) as conn:
            r = conn.execute(
                "SELECT is_default FROM tk_video_script_templates WHERE id=?",
                (template_id,),
            ).fetchone()
            if not r:
                return False
            if r["is_default"]:
                # Don't actually delete the default — archive instead.
                conn.execute(
                    "UPDATE tk_video_script_templates SET is_archived=1, is_default=0 WHERE id=?",
                    (template_id,),
                )
            else:
                conn.execute(
                    "DELETE FROM tk_video_script_templates WHERE id=?",
                    (template_id,),
                )
            conn.commit()
            return True

    # ------------------------------------------------------------------
    # Per-video script generation
    # ------------------------------------------------------------------

    def _gather_video_context(self, video_id: int) -> dict[str, Any]:
        """Pull pack info needed for prompt building."""
        with _open_db(self.db_path) as conn:
            row = conn.execute(
                """SELECT v.id AS video_id, v.pack_id, v.caption AS video_caption,
                          p.display_name, p.total_stickers, p.pack_uid, p.series_id,
                          s.style_anchor, s.palette, s.pack_archetype, s.metadata_json
                     FROM tk_videos v
                LEFT JOIN packs p        ON p.id = v.pack_id
                LEFT JOIN pack_series s  ON s.id = p.series_id
                    WHERE v.id = ?""",
                (video_id,),
            ).fetchone()
            if not row:
                raise ValueError(f"video #{video_id} not found")
            ctx = dict(row)
            sk = conn.execute(
                "SELECT seller_sku FROM tkshop_products WHERE pack_id=? ORDER BY id DESC LIMIT 1",
                (row["pack_id"],),
            ).fetchone()
            ctx["seller_sku"] = (sk["seller_sku"] if sk else "") or ""

        briefs: list[str] = []
        try:
            md = json.loads(ctx.get("metadata_json") or "{}")
            for pb in md.get("preview_briefs", [])[:3]:
                briefs.extend(pb.get("stickers", [])[:5])
        except Exception:
            pass
        ctx["briefs_sample"] = briefs
        return ctx

    def generate_scripts(
        self,
        video_id: int,
        *,
        template_id: int,
        variants: list[str] = ("A",),
        model: Optional[str] = None,
    ) -> dict[str, Any]:
        """Generate 1+ script variants for a video. Each variant is a
        complete script row. Variants run in parallel (text gen is cheap).
        Returns ``{script_ids, errors}``.
        """
        template = self.get_template(template_id)
        if not template:
            raise ValueError(f"template #{template_id} not found")
        ctx = self._gather_video_context(video_id)
        chosen_model = model or os.getenv(
            "TKVIDEO_SCRIPT_MODEL", os.getenv("OPENAI_MODEL", "gpt-5.4"),
        )

        def _one(variant: str) -> dict[str, Any]:
            prompt = build_script_main_prompt(
                template_name=template["name"],
                music_style=template.get("music_style") or "",
                pack_display_name=ctx.get("display_name") or "",
                pack_archetype=ctx.get("pack_archetype") or "",
                style_anchor=ctx.get("style_anchor") or "",
                palette=ctx.get("palette") or "",
                total_stickers=int(ctx.get("total_stickers") or 0),
                sticker_briefs_sample=ctx.get("briefs_sample") or [],
                seller_sku=ctx.get("seller_sku") or "",
                scene_blueprint=template["scene_blueprint"],
                variant_label=variant,
            )
            try:
                main_text = self.router.text_complete(
                    prompt,
                    model=chosen_model,
                    system=SCRIPT_SYSTEM_PROMPT,
                    temperature=0.75,
                    task=f"video_script:main:{variant}",
                    related_table="tk_videos",
                    related_id=video_id,
                )
                payload = self.router.extract_json(
                    main_text,
                    schema=SCRIPT_EXTRACT_SCHEMA,
                    instructions=SCRIPT_EXTRACT_INSTRUCTIONS,
                    model=chosen_model,
                    max_retries=1,
                    task=f"video_script:extract:{variant}",
                    related_table="tk_videos",
                    related_id=video_id,
                )
            except Exception as e:  # noqa: BLE001
                logger.exception("video script gen failed for video #%d variant %s",
                                 video_id, variant)
                return {"variant": variant, "ok": False,
                        "error": f"{type(e).__name__}: {e}"}
            return {
                "variant": variant, "ok": True,
                "main_text": main_text, "payload": payload,
            }

        results: list[dict[str, Any]] = []
        with ThreadPoolExecutor(max_workers=min(3, len(variants))) as ex:
            results = list(ex.map(_one, list(variants)))

        # Persist successful ones
        n = _now()
        script_ids: list[int] = []
        errors: list[dict] = []
        with _open_db(self.db_path) as conn:
            for r in results:
                if not r.get("ok"):
                    errors.append({"variant": r["variant"], "error": r.get("error")})
                    continue
                p = r["payload"] or {}
                ai_scenes = p.get("scenes") or []
                # Merge AI-filled fields into the blueprint to produce final scenes
                merged_scenes: list[dict] = []
                for i, blueprint_scene in enumerate(template["scene_blueprint"], start=1):
                    ai_match = next(
                        (a for a in ai_scenes if int(a.get("idx") or 0) == i),
                        None,
                    )
                    merged_scenes.append({
                        "idx": i,
                        "kind": blueprint_scene.get("kind") or "",
                        "label": blueprint_scene.get("label") or "",
                        "duration_s": int(blueprint_scene.get("duration_s") or 0),
                        "intent": blueprint_scene.get("intent") or "",
                        "visual_brief_template": blueprint_scene.get("visual_brief") or "",
                        "caption": (ai_match.get("caption") if ai_match else "") or "",
                        "b_roll_brief": (ai_match.get("b_roll_brief") if ai_match else "") or "",
                        "hashtags": list((ai_match.get("hashtags") if ai_match else []) or []),
                    })
                total_dur = sum(s["duration_s"] for s in merged_scenes)
                cur = conn.execute(
                    """INSERT INTO tk_video_scripts
                        (video_id, template_id, template_name, variant_label,
                         status, total_duration_s, music_style,
                         music_suggestions, scenes_json, ai_model,
                         raw_main_text, created_at, updated_at)
                       VALUES (?, ?, ?, ?, 'draft', ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (
                        video_id, template_id, template["name"], r["variant"],
                        total_dur, template.get("music_style") or "",
                        json.dumps(p.get("music_suggestions") or [], ensure_ascii=False),
                        json.dumps(merged_scenes, ensure_ascii=False),
                        chosen_model, r["main_text"],
                        n, n,
                    ),
                )
                script_ids.append(cur.lastrowid)
            conn.commit()
        return {"script_ids": script_ids, "errors": errors,
                "model": chosen_model}

    # ------------------------------------------------------------------
    # Read / list / edit
    # ------------------------------------------------------------------

    def list_scripts(self, video_id: int) -> list[dict]:
        with _open_db(self.db_path) as conn:
            rows = conn.execute(
                """SELECT id, video_id, template_id, template_name, variant_label,
                          status, total_duration_s, music_style, music_suggestions,
                          scenes_json, ai_model, created_at, updated_at
                     FROM tk_video_scripts
                    WHERE video_id = ?
                    ORDER BY id DESC""",
                (video_id,),
            ).fetchall()
        out = []
        for r in rows:
            d = dict(r)
            try:
                d["scenes"] = json.loads(d.get("scenes_json") or "[]")
            except Exception:
                d["scenes"] = []
            try:
                d["music_suggestions"] = json.loads(d.get("music_suggestions") or "[]")
            except Exception:
                d["music_suggestions"] = []
            out.append(d)
        return out

    def get_script(self, script_id: int) -> Optional[dict]:
        with _open_db(self.db_path) as conn:
            r = conn.execute(
                "SELECT * FROM tk_video_scripts WHERE id=?", (script_id,),
            ).fetchone()
        if not r:
            return None
        d = dict(r)
        try:
            d["scenes"] = json.loads(d.get("scenes_json") or "[]")
        except Exception:
            d["scenes"] = []
        try:
            d["music_suggestions"] = json.loads(d.get("music_suggestions") or "[]")
        except Exception:
            d["music_suggestions"] = []
        return d

    def update_scene(
        self,
        script_id: int,
        scene_idx: int,
        *,
        caption: Optional[str] = None,
        b_roll_brief: Optional[str] = None,
        duration_s: Optional[int] = None,
        hashtags: Optional[list[str]] = None,
    ) -> bool:
        s = self.get_script(script_id)
        if not s:
            return False
        scenes = s.get("scenes") or []
        for sc in scenes:
            if int(sc.get("idx") or 0) != int(scene_idx):
                continue
            if caption is not None:
                sc["caption"] = caption[:200]
            if b_roll_brief is not None:
                sc["b_roll_brief"] = b_roll_brief
            if duration_s is not None:
                sc["duration_s"] = max(1, int(duration_s))
            if hashtags is not None:
                sc["hashtags"] = [str(h).lstrip("#").strip() for h in hashtags if str(h).strip()]
            break
        else:
            return False
        total_dur = sum(int(sc.get("duration_s") or 0) for sc in scenes)
        with _open_db(self.db_path) as conn:
            conn.execute(
                """UPDATE tk_video_scripts
                      SET scenes_json=?, total_duration_s=?, updated_at=?
                    WHERE id=?""",
                (json.dumps(scenes, ensure_ascii=False), total_dur, _now(), script_id),
            )
            conn.commit()
        return True

    def update_script_status(self, script_id: int, status: str) -> bool:
        if status not in ("draft", "approved", "archived"):
            raise ValueError(f"invalid status: {status}")
        with _open_db(self.db_path) as conn:
            cur = conn.execute(
                "UPDATE tk_video_scripts SET status=?, updated_at=? WHERE id=?",
                (status, _now(), script_id),
            )
            conn.commit()
            return cur.rowcount > 0

    def delete_script(self, script_id: int) -> bool:
        with _open_db(self.db_path) as conn:
            cur = conn.execute("DELETE FROM tk_video_scripts WHERE id=?", (script_id,))
            conn.commit()
            return cur.rowcount > 0

    # ------------------------------------------------------------------
    # DOCX export
    # ------------------------------------------------------------------

    def export_script_docx(self, script_id: int) -> bytes:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        s = self.get_script(script_id)
        if not s:
            raise ValueError(f"script #{script_id} not found")

        # Pull pack name for header
        with _open_db(self.db_path) as conn:
            row = conn.execute(
                """SELECT v.id, p.display_name, p.pack_uid
                     FROM tk_videos v
                LEFT JOIN packs p ON p.id = v.pack_id
                    WHERE v.id = ?""",
                (s["video_id"],),
            ).fetchone()
        pack_name = (row["display_name"] if row else "") or "(pack)"

        doc = Document()
        # Title
        h = doc.add_heading(f"TikTok 视频脚本 — {pack_name}", level=0)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

        # Meta
        meta = doc.add_paragraph()
        meta.add_run("模板：").bold = True
        meta.add_run(f"{s.get('template_name') or '—'}    ")
        meta.add_run("变体：").bold = True
        meta.add_run(f"{s.get('variant_label') or 'A'}    ")
        meta.add_run("总时长：").bold = True
        meta.add_run(f"{s.get('total_duration_s') or 0} 秒    ")
        meta.add_run("状态：").bold = True
        meta.add_run(f"{s.get('status') or 'draft'}")

        # Music
        doc.add_heading("背景音乐建议", level=1)
        ms_para = doc.add_paragraph()
        ms_para.add_run(f"风格：{s.get('music_style') or '(未指定)'}\n").bold = False
        for m in (s.get("music_suggestions") or []):
            line = doc.add_paragraph(style="List Bullet")
            line.add_run(f"{m.get('name','')}").bold = True
            extras = " | ".join(filter(None, [m.get("mood"), m.get("bpm")]))
            if extras:
                line.add_run(f" — {extras}")
            note = m.get("note") or ""
            if note:
                line.add_run(f"\n   说明：{note}")

        # Scenes
        doc.add_heading("分场脚本", level=1)
        scenes = s.get("scenes") or []
        for sc in scenes:
            head = doc.add_heading(
                f"场景 {sc.get('idx')} · {sc.get('label') or sc.get('kind')} "
                f"（{sc.get('duration_s') or 0} 秒）",
                level=2,
            )

            t = doc.add_table(rows=4, cols=2)
            t.style = "Light List Accent 1"
            t.autofit = False
            for i, w in enumerate([Cm(3.0), Cm(13.0)]):
                t.columns[i].width = w
            rows_kv = [
                ("意图", sc.get("intent") or ""),
                ("画面建议（操作员）", sc.get("b_roll_brief") or sc.get("visual_brief_template") or ""),
                ("屏幕字幕（caption）", sc.get("caption") or ""),
                ("Hashtags", " ".join(f"#{h}" for h in (sc.get("hashtags") or []))),
            ]
            for i, (k, v) in enumerate(rows_kv):
                cell_k = t.cell(i, 0)
                cell_v = t.cell(i, 1)
                cell_k.text = k
                for r in cell_k.paragraphs[0].runs:
                    r.bold = True
                cell_v.text = v
            doc.add_paragraph()  # spacer

        # Footer / notes
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ts = time.strftime("%Y-%m-%d %H:%M", time.localtime(s.get("updated_at") or _now()))
        footer.add_run(f"导出时间：{ts}    脚本 ID：#{s['id']}").italic = True

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ---------------------------------------------------------------------------
# Module singleton
# ---------------------------------------------------------------------------

_svc: Optional[VideoScriptService] = None


def get_video_script_service() -> VideoScriptService:
    global _svc
    if _svc is None:
        _svc = VideoScriptService()
    return _svc
